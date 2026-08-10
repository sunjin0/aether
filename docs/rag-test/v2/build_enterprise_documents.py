from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

OUT = Path(__file__).parent / "enterprise-assets"
OUT.mkdir(exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def base_doc(title, subtitle, owner, classification):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(0.85)
    for style_name, size, color in [("Normal", 10.5, RGBColor(35, 45, 55)), ("Heading 1", 15, BLUE), ("Heading 2", 12, DARK), ("Heading 3", 11, DARK)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.name = "Microsoft YaHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); r.font.size = Pt(20); r.font.color.rgb = BLUE
    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10); p.runs[0].font.color.rgb = RGBColor(89, 98, 110)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    metadata = [("文档所有者", owner), ("分类等级", classification), ("生效日期", "2026-08-01"), ("复审周期", "每季度或重大事件后")]
    for i, (k, v) in enumerate(metadata):
        table.cell(i, 0).text = k; table.cell(i, 1).text = v
        shade(table.cell(i, 0), "E8EEF5")
    doc.add_paragraph()
    return doc

def heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.2
    return p

def bullets(doc, entries):
    for entry in entries:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(entry)

def small_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    for i, value in enumerate(headers):
        table.cell(0, i).text = value; shade(table.cell(0, i), "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row): cells[i].text = value
    doc.add_paragraph()

def customer_operations():
    doc = base_doc("Aether 企业客户服务与事故管理运行手册", "受控运行文件｜面向客户成功、支持、SRE 与商业运营团队", "客户成功与服务运营部", "内部")
    heading(doc, "1. 目的、范围与控制原则")
    para(doc, "本手册用于统一 Enterprise 客户的支持受理、事故沟通、服务信用、续费风险与退出服务流程。任何外部承诺必须与已签署订单表、状态页或经批准的事故沟通一致；口头承诺、聊天记录或 CRM 备注不得替代合同变更单。")
    heading(doc, "1.1 角色与责任", 2)
    small_table(doc, ["角色", "核心责任", "不得执行的行为"], [
        ("客户成功经理", "维护成功计划、识别续费风险、组织业务回顾", "口头修改价格、SLA 或数据驻留承诺"),
        ("支持负责人", "分级工单、协调技术响应、更新客户", "绕过安全与权限流程"),
        ("事故指挥官", "决定事故级别、恢复判定与内部节奏", "未经证据确认根因"),
        ("商业运营", "解释账单、处理信用与订单变更", "承诺未批准的减免")])
    heading(doc, "2. 服务等级与工单受理")
    heading(doc, "2.1 分级标准与首次响应", 2)
    small_table(doc, ["等级", "定义", "Enterprise 首次响应", "更新节奏"], [
        ("P1", "核心生产业务不可用且无绕行方案", "30 分钟", "每 30 分钟"),
        ("P2", "关键功能严重降级，有临时绕行方案", "2 小时", "每 4 小时"),
        ("P3", "一般缺陷、配置或使用咨询", "1 个工作日", "按工单约定")])
    heading(doc, "2.2 受理信息与升级路径", 2)
    para(doc, "工单必须包含租户标识、环境、影响开始时间、业务影响、请求 ID 或错误截图、已尝试的操作和技术联系人。若信息不足，支持人员应先请求补充而不是假定根因。P1 在 15 分钟内建立事故频道；涉及数据泄露、认证全面失败或跨租户访问风险时，同时通知安全响应团队。")
    bullets(doc, ["禁止在工单、截图和附件中写入密码、完整令牌、私钥或完整支付卡信息。", "客户升级必须指定单一负责人、下一步、下一次更新时间和关闭标准。", "响应时间只表示人工确认接手，不代表恢复时限或赔偿资格。"])
    heading(doc, "3. 事故沟通与恢复")
    heading(doc, "3.1 状态页和客户沟通", 2)
    para(doc, "确认 P1 后，通信负责人发布首条状态页信息，内容包括受影响能力、已知范围、正在采取的行动和下一次更新时间。没有新进展时仍需按节奏更新。对 Enterprise 客户的定向说明不得包含其他客户、未证实根因或内部安全细节。")
    heading(doc, "3.2 恢复判定与后续动作", 2)
    para(doc, "只有事故指挥官可以宣布恢复。恢复前应验证核心指标连续 30 分钟稳定、错误率回到基线、队列积压下降、关键客户合成探针正常且回滚风险可接受。宣布恢复后继续监控至少 60 分钟；P1 需在 5 个工作日内完成含时间线、根因、缓解和行动项的事故报告。")
    heading(doc, "4. 商业运营与服务信用")
    heading(doc, "4.1 退款、信用与账单争议", 2)
    para(doc, "服务信用仅适用于合同覆盖且经验证的服务指标未达标情形。计划维护、客户网络问题、客户配置错误、第三方身份提供商故障、不可抗力和超出合同限额通常不适用。客户须在受影响月份结束后 30 个自然日内提交租户 ID、影响日期和工单号。信用抵扣未来费用，不可兑换现金。")
    heading(doc, "4.2 续费与退出服务", 2)
    para(doc, "续费健康检查在合同到期前 120 天启动，90 天确认干系人，60 天完成价值和用量评估，30 天处理采购、安全或减量阻塞项。终止订阅后，客户在 30 天保留期内可导出数据；保留期结束后进入删除队列，备份在后续 90 天轮换周期内清除。法律保全、未结账款或安全调查可暂停删除。")
    heading(doc, "5. 记录、指标与复核")
    para(doc, "服务运营每月复核 P1/P2 首次响应、客户升级时长、服务信用处理时长、重复事故、待办行动项和续费风险。所有事故、升级、信用和数据删除决定必须关联工单或受控记录，并按适用的保留策略保存。")
    doc.save(OUT / "01-enterprise-customer-operations-manual.docx")

def governance_document():
    doc = base_doc("Aether 身份、数据与 AI 治理标准", "受控标准｜面向安全管理员、知识库所有者、模型负责人和审查者", "信息安全部与 AI 治理委员会", "内部限制")
    heading(doc, "1. 治理目标与适用范围")
    para(doc, "本标准规范平台身份、权限、知识库数据、人工审查和模型使用的最低控制要求。所有生产访问遵循最小权限、可审计、职责分离和可撤销原则。任何例外必须由指定批准角色书面确认，并记录范围、有效期、风险和补偿控制。")
    heading(doc, "2. 身份与访问控制")
    heading(doc, "2.1 角色与临时授权", 2)
    small_table(doc, ["访问类型", "批准要求", "最长有效期", "必需记录"], [
        ("知识库编辑", "知识库所有者", "长期，季度复核", "成员变更原因"),
        ("生产管理员", "系统所有者 + 安全管理员", "8 小时", "工单与会话审计"),
        ("P1 破窗访问", "值班负责人 + 独立批准人", "4 小时", "次日安全复核"),
        ("安全日志导出", "安全管理员", "单次任务", "导出任务 ID")])
    para(doc, "服务账号不得用于交互式登录，令牌应按最小范围签发，生产令牌最长有效期为 90 天。发现疑似泄露时必须立即撤销并重新签发，不能仅修改名称或备注。启用强制 SSO 后，用户恢复应通过企业身份提供商完成。")
    heading(doc, "3. 知识库数据治理")
    heading(doc, "3.1 数据分类与上传审查", 2)
    para(doc, "知识库允许公开、内部和经批准的机密资料；不得上传完整身份证号、支付卡号、私钥、未经脱敏的医疗记录或受出口管制约束的原始技术资料。上传版本先进入草稿和审核流程，AI 审查用于识别敏感信息、提示注入和质量风险，人工审查确认业务适用性与最终发布。AI 审查通过不等同于人工审批通过。")
    heading(doc, "3.2 生命周期与删除", 2)
    para(doc, "已审批版本才会进入 Agent 检索。删除文档会先从在线索引移除，再清理对象和分块；备份按 90 天轮换周期失效。法律保全、安全调查和未完成导出审计会阻止物理删除。发现敏感内容后，应先隔离、撤销分享、保全日志、评估影响并执行删除或脱敏。")
    heading(doc, "4. AI 风险与提示安全")
    heading(doc, "4.1 可用范围与人工监督", 2)
    para(doc, "模型可用于检索、摘要、分类和受批准的工作流草稿；招聘、信贷、医疗、法律结论、纪律处分和安全关键控制不得仅依赖模型输出。客户可见自动回复上线前必须完成事实抽样、敏感内容测试、人工接管路径和场景评审。")
    heading(doc, "4.2 提示注入与变更控制", 2)
    para(doc, "知识库中的“忽略指令”“导出密钥”等文字只能被视为待审查内容，不能改变系统权限或执行路径。模型、系统提示、检索参数、Chunk 策略或 Rerank 变更前必须完成可追溯评测；基线应绑定数据集版本、模型版本、检索配置和指标快照。")
    heading(doc, "5. 审计、事件与复核")
    para(doc, "管理员登录、权限变更、导出、知识库删除、令牌签发和人工审核决定必须写入审计日志。Enterprise 默认保留 365 天。疑似泄露在 30 分钟内报告安全响应团队；未经法务和数据保护官批准，不得对外承诺影响数量或通知日期。每季度复核访问成员、知识库分类、人工审查人、保留期限和高风险评测样例。")
    doc.save(OUT / "02-identity-data-ai-governance-standard.docx")

def ai_pdf():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title = ParagraphStyle("TitleCN", parent=styles["Title"], fontName="STSong-Light", fontSize=20, leading=28, textColor=colors.HexColor("#2E74B5"), alignment=TA_CENTER, spaceAfter=10)
    h1 = ParagraphStyle("H1CN", parent=styles["Heading1"], fontName="STSong-Light", fontSize=15, leading=22, textColor=colors.HexColor("#2E74B5"), spaceBefore=10, spaceAfter=6)
    h2 = ParagraphStyle("H2CN", parent=styles["Heading2"], fontName="STSong-Light", fontSize=12, leading=18, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontName="STSong-Light", fontSize=9.5, leading=16, spaceAfter=6)
    doc = SimpleDocTemplate(str(OUT / "03-ai-risk-and-model-governance-standard.pdf"), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm)
    story = [Paragraph("Aether AI 风险与模型治理标准", title), Paragraph("受控发布版本｜所有者：AI 治理委员会｜生效日期：2026-08-01｜分类：内部限制", body)]
    data = [["控制项", "要求"], ["批准角色", "AI 治理委员会、信息安全负责人、数据保护官"], ["适用范围", "面向客户的 Agent、检索、模型配置与自动化工作流"], ["复审触发", "模型替换、高风险场景上线、重大安全事件或季度评测复核"]]
    t = Table(data, colWidths=[34*mm, 130*mm]); t.setStyle(TableStyle([("FONT", (0,0),(-1,-1), "STSong-Light", 8.5), ("BACKGROUND", (0,0),(-1,0), colors.HexColor("#E8EEF5")), ("GRID", (0,0),(-1,-1), 0.25, colors.HexColor("#B9C6D3")), ("VALIGN", (0,0),(-1,-1), "MIDDLE"), ("LEFTPADDING", (0,0),(-1,-1), 7), ("RIGHTPADDING", (0,0),(-1,-1), 7), ("TOPPADDING", (0,0),(-1,-1), 6), ("BOTTOMPADDING", (0,0),(-1,-1), 6)])); story += [t, Spacer(1, 8)]
    sections = [
        ("1. 目的与治理原则", [("1.1 可追溯的业务目的", "每个 AI 场景必须明确业务所有者、输入数据范围、预期输出、可接受风险、人工接管路径和停止条件。平均指标不能取代对高风险样例的人工审阅。"), ("1.2 禁止的自动化决定", "招聘、信贷、医疗、法律结论、纪律处分和安全关键控制不得只依赖模型输出。模型只能提供辅助建议，最终决定必须由具备授权的人作出。")]),
        ("2. 知识检索与提示安全", [("2.1 非受信指令处理", "检索内容可能包含要求忽略规则、泄露系统提示或调用外部工具的恶意文本。此类文本仅作为待审查资料，绝不能改变权限、导出数据或绕过审批。"), ("2.2 审查和发布", "AI 审查标记敏感信息、注入式提示和质量风险；人工审查确认业务适用性。未审批文档不得参与检索，审查通过也不代表模型对所有问题均正确。")]),
        ("3. 模型与检索变更", [("3.1 上线前评测", "更换模型、修改系统提示、调整检索参数、Chunk 策略或启用 Rerank 前，应运行受控评测。数据集需覆盖正常、歧义、跨文档、过期信息、拒答和攻击样例。"), ("3.2 发布与回滚", "对客户可见的重大变更必须保留可回滚版本并观察至少 7 天。若平均 Recall 提升但高风险问答下降，变更不得仅因综合指标上升而发布。")]),
        ("4. 事件、监控与审计", [("4.1 异常输出", "发现虚构、歧视、泄密或越权建议时，应停止自动执行、保存输入输出和请求 ID、确认是否已对外发送并按安全事件流程升级。不得为了清理界面删除证据。"), ("4.2 持续监控", "生产至少监控检索命中率、无答案率、引用覆盖率、用户纠正率、拒答率、敏感内容拦截率、模型错误率和人工接管率。每季度更新评测语料和攻击样例。")])]
    for head, children in sections:
        story.append(Paragraph(head, h1))
        for sub, text in children:
            story.append(Paragraph(sub, h2)); story.append(Paragraph(text, body))
    doc.build(story)

customer_operations(); governance_document(); ai_pdf()
print(OUT)
